import os
import json
import re
import asyncio
from pathlib import Path
from typing import List, Optional

from dotenv import load_dotenv
from pydantic import BaseModel, Field, ValidationError

from docx import Document

from autogen_agentchat.agents import AssistantAgent
from autogen_ext.models.openai import OpenAIChatCompletionClient


# ============================================================
# 1. 加载环境变量
# ============================================================

load_dotenv("config.env")

DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY")

if not DEEPSEEK_API_KEY:
    raise RuntimeError(
        "未检测到 DEEPSEEK_API_KEY，请在 .env 文件中配置 DEEPSEEK_API_KEY。"
    )


# ============================================================
# 2. 定义结构化数据模型
# ============================================================

class Task(BaseModel):
    """
    单个待办任务
    """

    task_id: str = Field(
        description="任务编号，例如 TASK-001"
    )

    action: str = Field(
        description="任务内容"
    )

    owner: Optional[str] = Field(
        default=None,
        description="负责人。如果会议纪要中没有明确说明，则为 null"
    )

    deadline: Optional[str] = Field(
        default=None,
        description="截止时间。如果会议纪要中没有明确说明，则为 null"
    )

    priority: str = Field(
        default="普通",
        description="任务优先级，例如高、普通、低"
    )

    status: str = Field(
        default="待开始",
        description="任务状态，例如待开始、进行中、已完成、阻塞、待确认"
    )

    source: Optional[str] = Field(
        default=None,
        description="任务在会议纪要中的原始依据"
    )

    missing_info: List[str] = Field(
        default_factory=list,
        description="该任务缺失的信息，例如负责人、截止时间"
    )


class PendingConfirmation(BaseModel):
    """
    待人工确认事项
    """

    task_id: str = Field(
        description="需要确认的任务编号"
    )

    question: str = Field(
        description="需要人工确认的问题"
    )


class MeetingResult(BaseModel):
    """
    最终会议分析结果
    """

    meeting_title: Optional[str] = Field(
        default=None,
        description="会议主题"
    )

    meeting_date: Optional[str] = Field(
        default=None,
        description="会议日期"
    )

    participants: List[str] = Field(
        default_factory=list,
        description="参会人员列表"
    )

    summary: Optional[str] = Field(
        default=None,
        description="会议内容摘要"
    )

    tasks: List[Task] = Field(
        default_factory=list,
        description="会议待办任务列表"
    )

    missing_items: List[str] = Field(
        default_factory=list,
        description="会议整体缺失的信息"
    )

    review_notes: List[str] = Field(
        default_factory=list,
        description="审核 Agent 的审核说明"
    )

    pending_confirmation: List[PendingConfirmation] = Field(
        default_factory=list,
        description="需要人工确认的事项"
    )


# ============================================================
# 3. 初始化 DeepSeek 模型客户端
# ============================================================

model_client = OpenAIChatCompletionClient(
    model="deepseek-chat",
    api_key=DEEPSEEK_API_KEY,
    base_url="https://api.deepseek.com/v1",

    # DeepSeek 兼容 OpenAI API，
    # 但不是 AutoGen 内置模型名称，因此提供 model_info。
    model_info={
        "vision": False,
        "function_calling": True,
        "json_output": True,
        "family": "unknown",
    },
)


# ============================================================
# 4. 创建提取 Agent
# ============================================================

extract_agent = AssistantAgent(
    name="extract_agent",
    model_client=model_client,
    system_message="""
你是一个专业的会议纪要信息提取 Agent。

你的任务是：
1. 阅读会议纪要；
2. 提取会议主题；
3. 提取会议日期；
4. 提取参会人员；
5. 提取所有明确的待办事项；
6. 判断每个待办事项的负责人；
7. 判断每个待办事项的截止时间；
8. 判断每个任务的优先级；
9. 判断每个任务的当前状态；
10. 生成任务列表；
11. 找出负责人、截止时间等缺失信息；
12. 生成需要人工确认的问题。

重要规则：
- 只能根据会议纪要提取信息，不允许凭空猜测；
- 如果没有明确负责人，owner 必须为 null；
- 如果没有明确截止时间，deadline 必须为 null；
- “本周五”“月底前”“下周”等相对时间可以原样保留；
- 对每个任务记录原始依据 source；
- missing_info 中列出该任务缺少的信息；
- 如果负责人或截止时间缺失，任务状态可以设置为“待确认”；
- 如果会议纪要明确表示任务已经完成，status 设置为“已完成”；
- 如果会议纪要明确表示任务正在进行，status 设置为“进行中”；
- 如果会议纪要明确表示任务被阻塞，status 设置为“阻塞”；
- 如果没有状态信息，status 设置为“待开始”；
- 优先级只能根据会议纪要中的明确表达判断；
- 如果没有明确优先级，priority 设置为“普通”；
- 会议纪要中的普通讨论、意见、问题不一定是待办事项；
- 只有需要执行、完成、提交、准备、整理、开发、评估、修复、确认等动作时，才应提取为任务；
- 不允许将“下次会议讨论某事项”直接识别为当前待办任务；
- 参会人员必须按照会议纪要原文提取；
- 必须输出合法 JSON；
- 不要输出 Markdown；
- 不要输出解释文字。

输出格式必须严格如下：

{
  "meeting_title": "会议主题，没有则为 null",
  "meeting_date": "会议日期，没有则为 null",
  "participants": [
    "参会人员1",
    "参会人员2"
  ],
  "summary": "会议内容摘要",
  "tasks": [
    {
      "task_id": "TASK-001",
      "action": "任务内容",
      "owner": "负责人，没有则为 null",
      "deadline": "截止时间，没有则为 null",
      "priority": "高",
      "status": "待开始",
      "source": "会议纪要中的原始依据",
      "missing_info": []
    }
  ],
  "missing_items": [],
  "review_notes": [],
  "pending_confirmation": [
    {
      "task_id": "TASK-001",
      "question": "需要人工确认的问题"
    }
  ]
}
""",
)


# ============================================================
# 5. 创建审核 Agent
# ============================================================

review_agent = AssistantAgent(
    name="review_agent",
    model_client=model_client,
    system_message="""
你是一个严格的会议任务审核 Agent。

你的任务是审核提取 Agent 输出的会议任务列表。

审核要求：
1. 检查是否把普通讨论误识别为待办事项；
2. 检查每个任务是否包含明确的 action；
3. 检查负责人 owner 是否存在；
4. 检查截止时间 deadline 是否存在；
5. 检查 priority 是否合理；
6. 检查 status 是否合理；
7. 检查负责人是否属于参会人员名单；
8. 如果负责人不在参会人员名单中，必须在 review_notes 中标记异常；
9. 如果负责人或截止时间没有在原始纪要中明确出现，必须保留为 null；
10. 不允许根据常识、职位或上下文猜测负责人；
11. 不允许擅自补充会议纪要中不存在的截止时间；
12. 将任务缺失信息写入任务的 missing_info；
13. 将会议整体缺失信息写入 missing_items；
14. 对缺失负责人或截止时间的任务生成 pending_confirmation；
15. 如果任务缺失负责人或截止时间，必要时将 status 设置为“待确认”；
16. 给出简洁、明确的审核说明；
17. 必须保留原始会议纪要依据 source；
18. 必须输出合法 JSON；
19. 不要输出 Markdown；
20. 不要输出额外解释。

请基于“原始会议纪要”和“提取 Agent 结果”进行审核。

输出格式必须严格如下：

{
  "meeting_title": "会议主题，没有则为 null",
  "meeting_date": "会议日期，没有则为 null",
  "participants": [
    "参会人员1",
    "参会人员2"
  ],
  "summary": "会议内容摘要",
  "tasks": [
    {
      "task_id": "TASK-001",
      "action": "任务内容",
      "owner": "负责人，没有则为 null",
      "deadline": "截止时间，没有则为 null",
      "priority": "普通",
      "status": "待开始",
      "source": "原始纪要依据",
      "missing_info": []
    }
  ],
  "missing_items": [],
  "review_notes": [],
  "pending_confirmation": [
    {
      "task_id": "TASK-001",
      "question": "需要人工确认的问题"
    }
  ]
}
""",
)


# ============================================================
# 6. 读取会议纪要
# ============================================================

def read_docx(file_path: str) -> str:
    """
    读取 Word 文档中的段落和表格内容。
    """

    document = Document(file_path)

    contents = []

    # 读取普通段落
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if text:
            contents.append(text)

    # 读取 Word 表格
    for table in document.tables:
        for row in table.rows:
            cells = []

            for cell in row.cells:
                cell_text = cell.text.strip()

                if cell_text:
                    cells.append(cell_text)

            if cells:
                contents.append(" | ".join(cells))

    return "\n".join(contents).strip()


def read_minutes_auto(file_path: str) -> str:
    """
    根据文件扩展名自动读取会议纪要。

    支持：
    - .txt
    - .docx
    """

    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(
            f"会议纪要文件不存在：{file_path}"
        )

    suffix = path.suffix.lower()

    if suffix == ".txt":
        content = path.read_text(
            encoding="utf-8"
        ).strip()

    elif suffix == ".docx":
        content = read_docx(file_path)

    else:
        raise ValueError(
            f"暂不支持文件格式：{suffix}，目前仅支持 .txt 和 .docx"
        )

    if not content:
        raise ValueError("会议纪要内容为空")

    return content


# ============================================================
# 7. 处理模型返回结果
# ============================================================

def extract_json_from_text(text: str) -> dict:
    """
    从模型输出中提取 JSON。

    兼容以下情况：
    1. 模型直接输出 JSON；
    2. 模型输出 ```json ... ```；
    3. 模型在 JSON 前后输出少量说明文字。
    """

    text = text.strip()

    # 去除 Markdown 代码块
    text = re.sub(
        r"^```json\s*",
        "",
        text,
        flags=re.IGNORECASE
    )

    text = re.sub(
        r"^```\s*",
        "",
        text
    )

    text = re.sub(
        r"\s*```$",
        "",
        text
    )

    # 尝试直接解析
    try:
        data = json.loads(text)

        if not isinstance(data, dict):
            raise ValueError("模型返回的 JSON 不是对象")

        return data

    except json.JSONDecodeError:
        pass

    # 尝试截取第一个 JSON 对象
    start = text.find("{")
    end = text.rfind("}")

    if start != -1 and end != -1 and end > start:
        json_text = text[start:end + 1]
        data = json.loads(json_text)

        if not isinstance(data, dict):
            raise ValueError("模型返回的 JSON 不是对象")

        return data

    raise ValueError(
        f"无法从模型输出中解析 JSON：\n{text}"
    )


def get_last_message_text(task_result) -> str:
    """
    获取 AutoGen Agent 最后一条消息内容。
    """

    if not task_result.messages:
        raise ValueError("Agent 没有返回任何消息")

    last_message = task_result.messages[-1]
    content = last_message.content

    if isinstance(content, str):
        return content

    # 兼容某些消息内容为列表的情况
    if isinstance(content, list):
        text_parts = []

        for item in content:
            if isinstance(item, str):
                text_parts.append(item)

            elif isinstance(item, dict) and "text" in item:
                text_parts.append(item["text"])

        return "\n".join(text_parts)

    return str(content)


# ============================================================
# 8. 本地规则校验与人工确认事项生成
# ============================================================

def post_process_result(result: MeetingResult) -> MeetingResult:
    """
    对 Agent 输出结果进行本地规则校验和补充。

    主要功能：
    1. 检查负责人是否属于参会人员；
    2. 补充任务缺失信息；
    3. 生成 pending_confirmation；
    4. 对缺失关键字段的任务设置为“待确认”；
    5. 汇总会议整体缺失信息。
    """

    participant_set = {
        participant.strip()
        for participant in result.participants
        if participant and participant.strip()
    }

    existing_review_notes = list(result.review_notes)
    existing_missing_items = list(result.missing_items)

    pending_map = {
        item.task_id: item
        for item in result.pending_confirmation
    }

    for task in result.tasks:
        missing_info = list(task.missing_info)

        # ----------------------------------------------------
        # 检查 action
        # ----------------------------------------------------

        if not task.action or not task.action.strip():
            if "任务内容" not in missing_info:
                missing_info.append("任务内容")

        # ----------------------------------------------------
        # 检查负责人
        # ----------------------------------------------------

        if not task.owner:
            if "负责人" not in missing_info:
                missing_info.append("负责人")

            question = f"{task.action} 的负责人是谁？"

            pending_map[task.task_id] = PendingConfirmation(
                task_id=task.task_id,
                question=question
            )

        elif participant_set and task.owner not in participant_set:
            note = (
                f"{task.task_id} 的负责人“{task.owner}”"
                f"不在会议纪要提取到的参会人员名单中，请人工核实。"
            )

            if note not in existing_review_notes:
                existing_review_notes.append(note)

        # ----------------------------------------------------
        # 检查截止时间
        # ----------------------------------------------------

        if not task.deadline:
            if "截止时间" not in missing_info:
                missing_info.append("截止时间")

            question = f"{task.action} 的截止时间是什么？"

            pending_map[task.task_id] = PendingConfirmation(
                task_id=task.task_id,
                question=question
            )

        # ----------------------------------------------------
        # 如果关键信息缺失，设置为待确认
        # ----------------------------------------------------

        if missing_info and task.status == "待开始":
            task.status = "待确认"

        task.missing_info = missing_info

        # ----------------------------------------------------
        # 汇总会议整体缺失信息
        # ----------------------------------------------------

        for item in missing_info:
            summary_item = f"{task.task_id} 缺少{item}"

            if summary_item not in existing_missing_items:
                existing_missing_items.append(summary_item)

    result.missing_items = existing_missing_items
    result.review_notes = existing_review_notes
    result.pending_confirmation = list(pending_map.values())

    return result


# ============================================================
# 9. 主流程
# ============================================================

async def run_meeting_assistant(
    minutes_file: str,
    output_file: str = "meeting_result.json"
) -> MeetingResult:
    """
    执行智能会议助手流程：

    1. 自动读取 TXT 或 DOCX 会议纪要；
    2. 提取 Agent 提取会议结构化信息；
    3. 审核 Agent 审核提取结果；
    4. 本地规则校验参与者、负责人、截止时间；
    5. 生成待人工确认事项；
    6. 输出结构化 JSON。
    """

    # --------------------------------------------------------
    # 步骤一：自动读取会议纪要
    # --------------------------------------------------------

    minutes = read_minutes_auto(minutes_file)

    print("正在读取会议纪要...")
    print(f"会议纪要文件：{minutes_file}")
    print(f"会议纪要长度：{len(minutes)} 字符")

    # --------------------------------------------------------
    # 步骤二：调用提取 Agent
    # --------------------------------------------------------

    extract_prompt = f"""
请分析下面的会议纪要，并提取会议结构化信息、待办事项、参会人员、
优先级、任务状态以及缺失信息。

【会议纪要】
{minutes}
"""

    print("\n正在调用提取 Agent...")

    extract_result = await extract_agent.run(
        task=extract_prompt
    )

    extract_text = get_last_message_text(extract_result)

    try:
        extracted_data = extract_json_from_text(extract_text)

    except Exception as exc:
        print("提取 Agent 返回内容如下：")
        print(extract_text)

        raise RuntimeError(
            "提取 Agent 的返回结果不是有效 JSON"
        ) from exc

    print("提取 Agent 已完成。")

    # --------------------------------------------------------
    # 步骤三：调用审核 Agent
    # --------------------------------------------------------

    review_prompt = f"""
请审核下面的会议纪要和提取 Agent 的结果。

【原始会议纪要】
{minutes}

【提取 Agent 结果】
{json.dumps(
    extracted_data,
    ensure_ascii=False,
    indent=2
)}

请输出审核后的最终结构化 JSON 结果。
"""

    print("\n正在调用审核 Agent...")

    review_result = await review_agent.run(
        task=review_prompt
    )

    review_text = get_last_message_text(review_result)

    try:
        reviewed_data = extract_json_from_text(review_text)

    except Exception as exc:
        print("审核 Agent 返回内容如下：")
        print(review_text)

        raise RuntimeError(
            "审核 Agent 的返回结果不是有效 JSON"
        ) from exc

    print("审核 Agent 已完成。")

    # --------------------------------------------------------
    # 步骤四：Pydantic 校验结构
    # --------------------------------------------------------

    try:
        final_result = MeetingResult.model_validate(
            reviewed_data
        )

    except ValidationError as exc:
        print("审核结果结构化校验失败：")
        print(exc)

        print("\n尝试使用提取 Agent 结果进行兜底...")

        try:
            final_result = MeetingResult.model_validate(
                extracted_data
            )

        except ValidationError as fallback_exc:
            print("提取 Agent 结果校验也失败：")
            print(fallback_exc)

            raise RuntimeError(
                "提取 Agent 和审核 Agent 的结果都无法通过数据模型校验"
            ) from exc

    # --------------------------------------------------------
    # 步骤五：本地规则校验和待确认事项生成
    # --------------------------------------------------------

    final_result = post_process_result(final_result)

    # --------------------------------------------------------
    # 步骤六：输出 JSON 文件
    # --------------------------------------------------------

    output_path = Path(output_file)

    output_path.write_text(
        json.dumps(
            final_result.model_dump(),
            ensure_ascii=False,
            indent=2
        ),
        encoding="utf-8"
    )

    print(
        f"\n处理完成，结构化结果已保存到："
        f"{output_path.absolute()}"
    )

    # 同时打印到终端
    print("\n最终结果：")

    print(
        json.dumps(
            final_result.model_dump(),
            ensure_ascii=False,
            indent=2
        )
    )

    return final_result


# ============================================================
# 10. 程序入口
# ============================================================

async def main():
    try:
        await run_meeting_assistant(
            minutes_file="meeting_minutes.docx",
            output_file="meeting_result.json"
        )

    finally:
        await model_client.close()


if __name__ == "__main__":
    asyncio.run(main())
